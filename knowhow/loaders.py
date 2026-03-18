"""Multi-format document loaders.

Supported formats:
- Text/Markdown (.txt, .md)
- Images (.png, .jpg, .jpeg, .gif, .webp) — via vision LLM
- PDF (.pdf) — via pymupdf or pdfplumber
- DOCX (.docx) — via python-docx
- CSV (.csv) — via stdlib csv
- HTML (.html, .htm) — via beautifulsoup4
- JSON (.json) — flattened key-value
"""

from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from knowhow.llm import LLMClient

# File extensions grouped by type
TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".log", ".text"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
CSV_EXTENSIONS = {".csv", ".tsv"}
HTML_EXTENSIONS = {".html", ".htm"}
JSON_EXTENSIONS = {".json", ".jsonl"}


def detect_file_type(path: str) -> str:
    """Detect file type from extension."""
    suffix = Path(path).suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return "text"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    if suffix in DOCX_EXTENSIONS:
        return "docx"
    if suffix in CSV_EXTENSIONS:
        return "csv"
    if suffix in HTML_EXTENSIONS:
        return "html"
    if suffix in JSON_EXTENSIONS:
        return "json"
    return "text"  # fallback: try as text


def load_file(path: str, llm: LLMClient | None = None) -> str:
    """Load any supported file and return its text content.

    Args:
        path: Path to the file.
        llm: LLM client (required for image files to use vision).
    """
    file_type = detect_file_type(path)

    if file_type == "text":
        return _load_text(path)
    elif file_type == "image":
        return _load_image(path, llm)
    elif file_type == "pdf":
        return _load_pdf(path)
    elif file_type == "docx":
        return _load_docx(path)
    elif file_type == "csv":
        return _load_csv(path)
    elif file_type == "html":
        return _load_html(path)
    elif file_type == "json":
        return _load_json(path)
    else:
        return _load_text(path)


def _load_text(path: str) -> str:
    return Path(path).read_text(encoding="utf-8")


def _load_image(path: str, llm: LLMClient | None = None) -> str:
    """Extract text/description from an image using vision LLM."""
    if llm is None:
        raise ValueError(
            f"Cannot load image '{path}' without an LLM client. "
            "Provide an LLM with vision support to extract text from images."
        )
    prompt = (
        "Extract ALL text from this image. If there is no text, describe the image "
        "content in detail including any diagrams, charts, tables, or visual information. "
        "Structure the output as clean, readable text."
    )
    return llm.describe_image(path, prompt)


def _load_pdf(path: str) -> str:
    """Load PDF file. Tries pymupdf first, then pdfplumber, then basic extraction."""
    # Try pymupdf (fitz)
    try:
        import fitz  # pymupdf
        doc = fitz.open(path)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        pass

    # Try pdfplumber
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        pass

    raise ImportError(
        "PDF support requires 'pymupdf' or 'pdfplumber'. "
        "Install with: pip install pymupdf  OR  pip install pdfplumber"
    )


def _load_docx(path: str) -> str:
    """Load DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "DOCX support requires 'python-docx'. "
            "Install with: pip install python-docx"
        )
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        paragraphs.append("\n".join(rows))
    return "\n\n".join(paragraphs)


def _load_csv(path: str) -> str:
    """Load CSV/TSV and convert to readable text."""
    suffix = Path(path).suffix.lower()
    delimiter = "\t" if suffix == ".tsv" else ","

    with open(path, encoding="utf-8") as f:
        reader = csv.DictReader(f, delimiter=delimiter)
        rows = []
        for i, row in enumerate(reader):
            parts = [f"{k}: {v}" for k, v in row.items() if v]
            rows.append(f"Row {i+1}: {', '.join(parts)}")
    return "\n".join(rows)


def _load_html(path: str) -> str:
    """Load HTML and extract text content."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "HTML support requires 'beautifulsoup4'. "
            "Install with: pip install beautifulsoup4"
        )
    html = Path(path).read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    # Remove script and style elements
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _load_json(path: str) -> str:
    """Load JSON/JSONL and convert to readable text."""
    content = Path(path).read_text(encoding="utf-8")

    # Try JSONL first
    if path.endswith(".jsonl"):
        lines = content.strip().split("\n")
        entries = []
        for i, line in enumerate(lines):
            try:
                obj = json.loads(line)
                entries.append(f"Entry {i+1}: {_flatten_json(obj)}")
            except json.JSONDecodeError:
                continue
        return "\n".join(entries)

    # Regular JSON
    data = json.loads(content)
    return _flatten_json(data)


def _flatten_json(obj: object, prefix: str = "") -> str:
    """Recursively flatten JSON into readable text."""
    if isinstance(obj, dict):
        parts = []
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else k
            parts.append(_flatten_json(v, key))
        return "\n".join(parts)
    elif isinstance(obj, list):
        parts = []
        for i, v in enumerate(obj):
            parts.append(_flatten_json(v, f"{prefix}[{i}]"))
        return "\n".join(parts)
    else:
        return f"{prefix}: {obj}" if prefix else str(obj)
